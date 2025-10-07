from docx import Document
import re

class DocxProcessor:
    def __init__(self):
        pass
    
    def _flush_buffer(self, buffer, sections, current_heading):
        """Push accumulated text into sections."""
        if buffer:
            text = "\n".join(buffer).strip()
            if text:
                sections.append({
                    "heading": current_heading,
                    "content": text,
                    "type": "text"
                })
            buffer.clear()

    def _table_to_markdown(self, table):
        """Convert DOCX table to Markdown format."""
        md_lines = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if i == 0:
                md_lines.append("| " + " | ".join(cells) + " |")
                md_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            else:
                md_lines.append("| " + " | ".join(cells) + " |")
        return "\n".join(md_lines).strip()

    def _handle_paragraph(self, para, buffer, sections, current_heading):
        """Process paragraph for heading or content."""
        text = para.text.strip()
        if not text:
            return current_heading  # skip empty

        style = para.style.name if para.style else ""

        # Detect heading
        if re.match(r"^Heading", style, re.I):
            self._flush_buffer(buffer, sections, current_heading)
            current_heading = text
        else:
            buffer.append(text)

        return current_heading

    def _handle_table(self, table, sections, current_heading, buffer):
        """Convert table and append to sections."""
        self._flush_buffer(buffer, sections, current_heading)
        md = self._table_to_markdown(table)
        if md:
            sections.append({
                "heading": current_heading,
                "content": md,
                "type": "table"
            })

    def process_file(self, file_path):
        """
        Parse DOCX and return structured sections with headings, text, and tables.
        Output: [{heading, content, type}]
        """
        doc = Document(file_path)
        sections, buffer = [], []
        current_heading = "Introduction"

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]
            if tag == "p":  # Paragraph
                para = next((p for p in doc.paragraphs if p._element == element), None)
                if para:
                    current_heading = self._handle_paragraph(para, buffer, sections, current_heading)

            elif tag == "tbl":  # Table
                table = next((t for t in doc.tables if t._element == element), None)
                if table:
                    self._handle_table(table, sections, current_heading, buffer)

        self._flush_buffer(buffer, sections, current_heading)
        return sections
